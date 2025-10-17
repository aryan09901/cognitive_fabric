import os
import json
import PyPDF2
import docx
from typing import List, Dict, Any, Optional
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

class FileProcessor:
    """
    File processor for handling various document formats
    """
    
    def __init__(self):
        self.supported_formats = {
            '.txt': self._process_text_file,
            '.pdf': self._process_pdf_file,
            '.docx': self._process_docx_file,
            '.json': self._process_json_file,
            '.md': self._process_markdown_file
        }
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """Process a file and extract its content"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        processor = self.supported_formats[file_extension]
        content = processor(file_path)
        
        return {
            'file_path': str(file_path),
            'file_name': file_path.name,
            'file_size': file_path.stat().st_size,
            'file_extension': file_extension,
            'content': content,
            'processed_at': os.path.getmtime(file_path)
        }
    
    def process_directory(self, directory_path: str) -> List[Dict[str, Any]]:
        """Process all supported files in a directory"""
        directory = Path(directory_path)
        
        if not directory.exists() or not directory.is_dir():
            raise ValueError(f"Directory not found: {directory_path}")
        
        processed_files = []
        
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                try:
                    processed_file = self.process_file(str(file_path))
                    processed_files.append(processed_file)
                    logger.info(f"Processed file: {file_path.name}")
                except Exception as e:
                    logger.error(f"Failed to process file {file_path}: {e}")
        
        return processed_files
    
    def _process_text_file(self, file_path: Path) -> str:
        """Process text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            return content
        except UnicodeDecodeError:
            # Try different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                content = file.read()
            return content
    
    def _process_pdf_file(self, file_path: Path) -> str:
        """Process PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content = ""
                
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
                
                return content.strip()
        except Exception as e:
            logger.error(f"Failed to process PDF {file_path}: {e}")
            return f"Error processing PDF: {str(e)}"
    
    def _process_docx_file(self, file_path: Path) -> str:
        """Process DOCX file"""
        try:
            doc = docx.Document(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            return content.strip()
        except Exception as e:
            logger.error(f"Failed to process DOCX {file_path}: {e}")
            return f"Error processing DOCX: {str(e)}"
    
    def _process_json_file(self, file_path: Path) -> Dict[str, Any]:
        """Process JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = json.load(file)
            return content
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON in {file_path}: {e}")
            return {"error": f"Invalid JSON: {str(e)}"}
    
    def _process_markdown_file(self, file_path: Path) -> str:
        """Process Markdown file"""
        return self._process_text_file(file_path)
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from file"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        stat = file_path.stat()
        
        metadata = {
            'file_name': file_path.name,
            'file_size': stat.st_size,
            'file_extension': file_path.suffix.lower(),
            'created_time': stat.st_ctime,
            'modified_time': stat.st_mtime,
            'absolute_path': str(file_path.absolute())
        }
        
        # Add format-specific metadata
        if file_path.suffix.lower() == '.pdf':
            metadata.update(self._extract_pdf_metadata(file_path))
        elif file_path.suffix.lower() == '.docx':
            metadata.update(self._extract_docx_metadata(file_path))
        
        return metadata
    
    def _extract_pdf_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract PDF-specific metadata"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                info = pdf_reader.metadata
                
                return {
                    'pdf_author': getattr(info, 'author', None),
                    'pdf_title': getattr(info, 'title', None),
                    'pdf_subject': getattr(info, 'subject', None),
                    'pdf_creator': getattr(info, 'creator', None),
                    'pdf_producer': getattr(info, 'producer', None),
                    'pdf_page_count': len(pdf_reader.pages)
                }
        except Exception as e:
            logger.error(f"Failed to extract PDF metadata {file_path}: {e}")
            return {}
    
    def _extract_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract DOCX-specific metadata"""
        try:
            doc = docx.Document(file_path)
            core_properties = doc.core_properties
            
            return {
                'docx_author': core_properties.author,
                'docx_title': core_properties.title,
                'docx_subject': core_properties.subject,
                'docx_created': core_properties.created,
                'docx_modified': core_properties.modified,
                'docx_last_modified_by': core_properties.last_modified_by
            }
        except Exception as e:
            logger.error(f"Failed to extract DOCX metadata {file_path}: {e}")
            return {}
    
    def chunk_content(self, content: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
        """Split content into overlapping chunks"""
        if not content:
            return []
        
        chunks = []
        start = 0
        
        while start < len(content):
            end = start + chunk_size
            
            # If this isn't the last chunk, try to end at a sentence boundary
            if end < len(content):
                # Look for sentence endings near the chunk boundary
                sentence_endings = ['.', '!', '?', '\n\n']
                for i in range(min(50, len(content) - end)):
                    if content[end + i] in sentence_endings:
                        end = end + i + 1
                        break
            
            chunk = content[start:end].strip()
            if chunk:  # Only add non-empty chunks
                chunks.append(chunk)
            
            start = end - overlap  # Overlap with next chunk
        
        return chunks

# Global file processor instance
file_processor = FileProcessor()